# -*- coding: utf-8 -*-
"""
Created on Mon Dec  9 12:07:19 2024

@author: YimingL
"""

import json
import pickle
import os
import io
from docx import Document
from docx.shared import Inches, RGBColor


def save_dict_to_json(dict_list: dict, file_path: str):
    with open(file_path, 'w') as json_file:
        json.dump(dict_list, json_file, indent=4)
    print(f"Data has been saved to {file_path}")

def save_variable_to_pickle(variable, file_path: str):
    with open(file_path, 'wb') as file:
        pickle.dump(variable, file)
    print(f"Data has been saved as {file_path}")

def save_dataframe_to_csv(df, file_path):
    df.to_csv(file_path, index=False, encoding='utf-8-sig')
    print(f"Data has been saved to {file_path}")
    
def save_text_file(data, file_path):
    with open(file_path, 'w') as f_out:
        f_out.write(data)
    print(f"Data has been saved to {file_path}")
    

def save_plots_summarization_to_word(data_list: list, output_folder: str):
    """ 将 5 组体温图像和文字总结保存到 5 个 Word 文档中 """
    os.makedirs(output_folder, exist_ok=True)  # 创建输出文件夹（如果不存在）

    for i, group in enumerate(data_list):  # 遍历 5 组
        doc = Document()  # 创建 Word 文档
        doc.add_heading(f'Group {i+1} - Temperature Records', level=1)

        for j, item in enumerate(group):  # 遍历每组 8 个记录
            plot = item["human_reader_plt"]  # 获取图像对象
            rule_summary = item["rule_summarization"]
            gpt_summary = item["open_ai_summarization"]

            # 添加记录标题
            doc.add_heading(f"Record {j+1}", level=2)

            # 创建一个新段落
            p = doc.add_paragraph()

            # 插入 "Rule Summarization:" 并设置蓝色
            run_rule = p.add_run("Rule Summarization: ")
            run_rule.font.color.rgb = RGBColor(0, 0, 255)  # 蓝色

            # 插入规则总结文本（默认颜色）
            p.add_run(rule_summary + "\n")

            # 插入 "GPT-4 Summarization:" 并设置蓝色
            run_gpt = p.add_run("GPT-4 Summarization: ")
            run_gpt.font.color.rgb = RGBColor(0, 0, 255)  # 蓝色

            # 插入 GPT-4 总结文本（默认颜色）
            p.add_run(gpt_summary + "\n")

            # **保存图像到本地**
            image_path = os.path.join(output_folder, f"group_{i+1}_record_{j+1}.png")
            plot.savefig(image_path, dpi=150)  # 保存图片到本地
            
            # **将图片插入 Word**
            doc.add_picture(image_path, width=Inches(5))

        # **保存 Word 文档**
        output_path = os.path.join(output_folder, f"Group_{i+1}.docx")
        doc.save(output_path)
        print(f"✅ Saved: {output_path}")


