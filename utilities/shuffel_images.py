# -*- coding: utf-8 -*-
"""
Created on Thu Jan 23 20:11:26 2025

@author: Yiming Li
"""

import os
import random
from docx import Document
from docx.shared import Inches

def distribute_images(images_dir, num_groups=4):
    """
    将本地图像文件分组并随机打乱。
    Distribute local image files into groups and shuffle them.

    参数 (Parameters):
        images_dir (str): 存储图片的文件夹路径 (Directory containing images).
        num_groups (int): 分组数量 (Number of groups).
    
    返回 (Returns):
        list: 每个组的图片路径列表 (List of image paths for each group).
    """
    # 获取所有图片文件路径 (Get all image file paths)
    all_images = [os.path.join(images_dir, img) for img in os.listdir(images_dir) if img.endswith(('.png', '.jpg', '.jpeg'))]
    
    # 打乱图片顺序 (Shuffle images)
    random.shuffle(all_images)
    
    # 初始化分组 (Initialise groups)
    groups = [[] for _ in range(num_groups)]
    
    # 均匀分配图片到各组 (Evenly distribute images into groups)
    for i, image in enumerate(all_images):
        groups[i % num_groups].append(image)
    
    return groups

def create_word_docs(groups, output_dir):
    """
    为每个组生成一个 Word 文档，插入对应的图片。
    Create a Word document for each group and insert corresponding images.

    参数 (Parameters):
        groups (list): 分组后的图片路径列表 (List of image paths for each group).
        output_dir (str): 输出文档的文件夹路径 (Directory to save the Word documents).
    """
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    for i, group in enumerate(groups, start=1):
        doc = Document()
        doc.add_heading(f'Group {i}', level=1)

        for image_path in group:
            # 添加图片标题 (Add image title)
            #doc.add_paragraph(f"Image: {os.path.basename(image_path)}")
            # 插入图片 (Insert image)
            doc.add_picture(image_path, width=Inches(5))
            doc.add_paragraph("\n")  # 空行
            doc.add_paragraph("Please write summarization here: ")
            doc.add_paragraph("\n")  # 空行

        # 保存文档 (Save the document)
        doc.save(os.path.join(output_dir, f'Group_{i}.docx'))

# 主函数 (Main function)
if __name__ == "__main__":
    # 图片文件夹路径 (Path to the folder containing images)
    images_folder = "../data/record images/for_huma_reader_summarise"  # 替换为你的图片文件夹路径 Replace with your image folder path

    # 输出文档文件夹路径 (Path to save the output documents)
    output_folder = "../data/docs"  # 替换为保存文档的路径 Replace with your output folder path

    # 将图片分成 4 组 (Distribute images into 4 groups)
    grouped_images = distribute_images(images_folder, num_groups=4)

    # 为每组生成 Word 文档 (Create Word documents for each group)
    create_word_docs(grouped_images, output_folder)

    print(f"Documents saved to {output_folder}")
