# -*- coding: utf-8 -*-
"""
Created on Wed Mar  5 10:37:51 2025

@author: Yiming Li
"""
import pickle
import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, RGBColor


# vital_sign_sum = cases[0][0]['febrile']['febrile_rule_summarization']
with open('./data/cases_micro.pkl', 'rb') as pkl:
    cases = pickle.load(pkl)
# vital_sign_sum = cases[0][0]['febrile']['febrile_rule_summarization'] + '\n' + cases[0][0]['heart_rate']['hr_rule_summary'] + '\n' + cases[0][0]['systolic_blood_pressure']['sbp_rule_summary']

out_put_folder ={
    'febrile': './data/plot/febrile',
    'heart_rate': './data/plot/heart_rate',
    'systolic_blood_pressure': './data/plot/systolic_blood_pressure'
    }


def set_cell_font(cell, bold=False, color=None, font_name="Times New Roman"):
    """设置单元格字体格式（支持颜色和加粗）"""
    if cell.paragraphs:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = font_name
                run.font.size = Pt(10)
                run.font.bold = bold
                if color:
                    run.font.color.rgb = color
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_images_to_cell(cell, image_paths):
    """向单元格插入多张图片（垂直排列）"""
    cell.text = ""  # 清空单元格原有内容
    for img_path in image_paths:
        paragraph = cell.add_paragraph()
        run = paragraph.add_run()
        run.add_picture(img_path, width=Inches(5))  # 控制图片宽度
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 图片居中


def add_section(doc, data):
    """添加一个患者记录的表格（整合所有特殊处理）"""
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"

    # === 标题行处理 ===
    title_row = table.add_row()
    title_cell = title_row.cells[0]
    title_cell.merge(title_row.cells[1])
    title_cell.text = "INFECTIOUS DISEASES / MICROBIOLOGY REVIEW"
    set_cell_font(title_cell, bold=True)

    # === 数据行处理 ===
    for key, value in data.items():
        row = table.add_row()
        left_cell = row.cells[0]
        right_cell = row.cells[1]

        # 左侧标题处理（保持默认黑色）
        left_cell.text = key
        set_cell_font(left_cell, bold=True)

        # === 特殊栏位处理 ===
        if key == "Imaging":  # Imaging栏图片处理
            if isinstance(value, list) and len(value) > 0:
                add_images_to_cell(right_cell, value)
            elif value and value != "-":
                add_images_to_cell(right_cell, [value])
            else:
                right_cell.text = "-"
                set_cell_font(right_cell)

        elif key == "Vital signs rules" or key == "Vital signs ds":  # Background栏颜色处理
            right_cell.text = ""
            lines = value.split("\n")
            for line in lines:
                paragraph = right_cell.add_paragraph()
                run = paragraph.add_run(line)
                run.font.name = "Times New Roman"
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0, 0, 255)  # 蓝色

        else:  # 其他常规栏位
            right_cell.text = str(value)
            set_cell_font(right_cell)

    doc.add_paragraph("\n")  # 表格间空行

# 定义数据（Imaging 字段改为图片路径列表）

for i, case_ in enumerate(cases):
    for j, example in enumerate(case_):
        for key in out_put_folder.keys():
            os.makedirs(out_put_folder[key], exist_ok=True)
            example[key]['human_reader_plt'].savefig(os.path.join(out_put_folder[key], f"group_{i+1}_record_{j+1}.png"), dpi=150)  # 保存图片到本地
        rule_sum = example['febrile']['febrile_rule_summarization'] + '\n' + example['cv_hmd_rule_sum']
        ds_sum = example['ds_summary_all']
        micro_sum = example['micro_summary']
        
        patient1 = {
            "Consult reason": "Positive blood cultures",
            "*Diagnosis**": "E coli bacteraemia, presumed urinary source",
            "Background": "Diabetes\nHypertension\nBPH -- long term catheter",
            "Presentation & Clinical Progress": "Admitted 26/2 with lower abdo pain and fever",
            "Antibiotic history": "26/2 → co-amox iv",
            "Vital signs rules": rule_sum,
            "Vital signs ds": ds_sum,
            "Physical exam": "Abdo soft and non tender",
            "Micro results": micro_sum,
            "Blood results": "28/2 WCC 14 (falling, from 20 on 26/2); CRP 150 (from 300 on 26/2)",
            "Imaging": [  # 多张图片示例
                os.path.join(out_put_folder['febrile'], f"group_{i + 1}_record_{j + 1}.png"),
                os.path.join(out_put_folder['heart_rate'], f"group_{i + 1}_record_{j + 1}.png"),
                os.path.join(out_put_folder['systolic_blood_pressure'], f"group_{i + 1}_record_{j + 1}.png")
            ],
            "Discussion": "With team F1, David",
            "**Advice**": "Continue current antibiotics\nConsider changing catheter while on antibiotics\nIf remains febrile consider ultrasound of urinary tract\nWe will review with final microbiology results",
            "Signed": "David Eyre, Consultant in Infection",
            "Responsible consultant": "David Eyre"
        }
    
    # 生成文档
        doc = Document()
        add_section(doc, patient1)
        output_path = os.path.join('./data/patient_review/', f"Group_{i + 1}_{j + 1}.docx")
        doc.save(output_path)